"""Tests for the workflow project type end-to-end:

- meta-init accepts --project-type workflow
- decision_log.workflow_config is required and validated
- scaffold emits inbox/outbox/state/kb_brief/orchestrator + run_workflow.py
- Stage 4 finalize check accepts a registry with tracked_doc/comment_reply outputs
- run-workflow CLI invokes the scaffold's runner against an input docx
"""
from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path

import pytest
import yaml

from meta_compiler.artifacts import build_paths, ensure_layout, save_manifest
from meta_compiler.project_types import VALID_PROJECT_TYPES, project_type_choices
from meta_compiler.stages.init_stage import run_meta_init
from meta_compiler.stages.scaffold_stage import run_scaffold
from meta_compiler.stages.workflow_stage import run_workflow
from meta_compiler.validation import validate_decision_log


def test_workflow_in_valid_project_types():
    assert "workflow" in VALID_PROJECT_TYPES
    assert "workflow" in project_type_choices()


def test_meta_init_accepts_workflow(tmp_path):
    workspace_root = tmp_path / "ws"
    artifacts_root = workspace_root / "workspace-artifacts"
    result = run_meta_init(
        workspace_root=workspace_root,
        artifacts_root=artifacts_root,
        project_name="Reviewer Replies",
        problem_domain="academic peer review",
        project_type="workflow",
    )
    paths = build_paths(artifacts_root)
    manifest = yaml.safe_load(paths.manifest_path.read_text(encoding="utf-8"))
    assert manifest["workspace_manifest"]["project_type"] == "workflow"
    # Wiki name is populated by later stages, not at init; verify the suffix mapping.
    from meta_compiler.artifacts import derive_wiki_name

    assert derive_wiki_name("Reviewer Replies", "workflow").endswith("Workflow Atlas")
    assert isinstance(result, dict)


def test_meta_init_rejects_unknown_project_type(tmp_path):
    workspace_root = tmp_path / "ws"
    artifacts_root = workspace_root / "workspace-artifacts"
    with pytest.raises(ValueError) as excinfo:
        run_meta_init(
            workspace_root=workspace_root,
            artifacts_root=artifacts_root,
            project_name="bad",
            problem_domain="bad",
            project_type="bogus",
        )
    assert "workflow" in str(excinfo.value)


# --- Decision-log validation --------------------------------------------------


def _minimal_workflow_decision_log(with_workflow_config: bool = True) -> dict:
    body = {
        "decision_log": {
            "meta": {
                "version": 1,
                "parent_version": 0,
                "reason_for_revision": "initial",
                "created": "2026-04-21T00:00:00+00:00",
                "project_name": "Reviewer Replies",
                "project_type": "workflow",
                "problem_statement_hash": "x" * 64,
                "wiki_version": "abc",
            },
            "conventions": [],
            "architecture": [],
            "scope": {"in_scope": [], "out_of_scope": []},
            "requirements": [],
            "open_items": [],
            "agents_needed": [],
        }
    }
    if with_workflow_config:
        body["decision_log"]["workflow_config"] = {
            "trigger": "inbox_watch",
            "inputs": [{"kind": "tracked_doc", "locator": "inbox/*.docx"}],
            "outputs": [
                {"kind": "tracked_edit", "target": "inbox/*.docx"},
                {"kind": "comment_reply", "target": "inline"},
            ],
            "state_keys": ["processed_comment_ids"],
            "escalation_policy": {"on_missing_kb_evidence": "pause_for_human"},
        }
    return body


def test_decision_log_requires_workflow_config_for_workflow_type():
    issues = validate_decision_log(_minimal_workflow_decision_log(with_workflow_config=False))
    assert any("workflow_config" in issue for issue in issues), issues


def test_decision_log_accepts_full_workflow_config():
    issues = validate_decision_log(_minimal_workflow_decision_log())
    assert all("workflow_config" not in i for i in issues), issues


def test_decision_log_rejects_invalid_trigger():
    body = _minimal_workflow_decision_log()
    body["decision_log"]["workflow_config"]["trigger"] = "fairy_dust"
    issues = validate_decision_log(body)
    assert any("trigger" in issue for issue in issues), issues


def test_decision_log_rejects_workflow_config_for_other_types():
    body = _minimal_workflow_decision_log()
    body["decision_log"]["meta"]["project_type"] = "report"
    # Make report-shaped (no code_architecture)
    issues = validate_decision_log(body)
    assert any("workflow_config" in issue for issue in issues), issues


# --- Scaffold layout ----------------------------------------------------------


def _seed_workflow_workspace(tmp_path: Path) -> tuple[Path, Path]:
    workspace_root = tmp_path / "ws"
    artifacts_root = workspace_root / "workspace-artifacts"
    paths = build_paths(artifacts_root)
    ensure_layout(paths)
    save_manifest(
        paths,
        {
            "workspace_manifest": {
                "project_type": "workflow",
                "name": "Reviewer Replies",
                "problem_domain": "test",
                "wiki": {"name": "Reviewer Replies Workflow Atlas", "version": "x"},
                "research": {"last_completed_stage": "2"},
                "decision_logs": [],
            }
        },
    )
    body = _minimal_workflow_decision_log()
    body["decision_log"]["agents_needed"] = [
        {
            "role": "tracked-edit-writer",
            "responsibility": "stub",
            "inputs": [{"name": "comment_reply", "modality": "document"}],
            "outputs": [{"name": "tracked_doc", "modality": "document"}],
            "key_constraints": [],
        }
    ]
    paths.decision_logs_dir.mkdir(parents=True, exist_ok=True)
    yaml_path = paths.decision_logs_dir / "decision_log_v1.yaml"
    yaml_path.write_text(yaml.safe_dump(body, sort_keys=False), encoding="utf-8")
    return workspace_root, artifacts_root


def test_scaffold_emits_workflow_layout(tmp_path):
    workspace_root, artifacts_root = _seed_workflow_workspace(tmp_path)
    paths = build_paths(artifacts_root)

    result = run_scaffold(
        artifacts_root=artifacts_root, decision_log_version=1
    )
    scaffold_root = Path(result["root"]) if "root" in result else paths.scaffolds_dir / "v1"
    assert (scaffold_root / "inbox").is_dir()
    assert (scaffold_root / "outbox").is_dir()
    assert (scaffold_root / "state").is_dir()
    assert (scaffold_root / "kb_brief").is_dir()
    assert (scaffold_root / "orchestrator" / "run_workflow.py").is_file()
    state = yaml.safe_load((scaffold_root / "state" / "state.yaml").read_text())
    assert state["workflow_state"]["schema_version"] == 1
    assert state["workflow_state"]["scaffold_version"] == 1


def test_scaffold_workflow_includes_canonical_agents(tmp_path):
    workspace_root, artifacts_root = _seed_workflow_workspace(tmp_path)
    paths = build_paths(artifacts_root)

    result = run_scaffold(artifacts_root=artifacts_root, decision_log_version=1)
    scaffold_root = paths.scaffolds_dir / "v1"
    agent_files = [p.stem for p in (scaffold_root / "agents").glob("*.md")]
    expected = {
        "workflow-conductor",
        "comment-reader",
        "kb-retriever",
        "response-author",
        "tracked-edit-writer",
        "kb-maintainer",
    }
    assert expected.issubset({a.lower().replace("_", "-") for a in agent_files}), (
        agent_files,
        expected,
    )


# --- run-workflow round-trip --------------------------------------------------


def _make_two_comments_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    p = doc.add_paragraph("Reviewer note one.")
    doc.add_comment([p.runs[0]], text="please clarify", author="R", initials="R")
    p2 = doc.add_paragraph("Reviewer note two.")
    doc.add_comment([p2.runs[0]], text="add citation", author="R", initials="R")
    doc.save(str(path))


def test_run_workflow_invokes_scaffold_orchestrator(tmp_path):
    workspace_root, artifacts_root = _seed_workflow_workspace(tmp_path)
    paths = build_paths(artifacts_root)

    run_scaffold(artifacts_root=artifacts_root, decision_log_version=1)

    docx_path = workspace_root / "input.docx"
    _make_two_comments_docx(docx_path)

    result = run_workflow(
        artifacts_root=artifacts_root,
        input_path=str(docx_path),
        task="reply-to-comments",
    )
    assert result["status"] == "ok", (result.get("stdout"), result.get("stderr"))
    assert result["scaffold_version"] == 1
    # The dispatch JSON is the last block of stdout. Find it.
    stdout = result["stdout"]
    json_start = stdout.find('{\n  "status": "dispatch_ready"')
    assert json_start != -1, stdout
    # Parse from json_start to end (the trailing payload).
    payload = json.loads(stdout[json_start:])
    # The dispatch JSON enumerates the scaffold's agent roles.
    assert payload["status"] == "dispatch_ready"
    assert any(
        agent and "tracked-edit-writer" in agent.lower()
        for agent in payload["plan"]["agents"]
    )


def test_run_workflow_rejects_non_workflow_scaffold(tmp_path):
    workspace_root = tmp_path / "ws"
    artifacts_root = workspace_root / "workspace-artifacts"
    paths = build_paths(artifacts_root)
    ensure_layout(paths)
    # Create a fake non-workflow scaffold v1
    scaffold_root = paths.scaffolds_dir / "v1"
    scaffold_root.mkdir(parents=True)
    (scaffold_root / "SCAFFOLD_MANIFEST.yaml").write_text(
        "scaffold:\n  project_type: report\n", encoding="utf-8"
    )
    (scaffold_root / "orchestrator").mkdir()
    (scaffold_root / "orchestrator" / "run_workflow.py").write_text("# stub\n")
    docx = workspace_root / "x.docx"
    docx.parent.mkdir(parents=True, exist_ok=True)
    docx.write_text("not really a docx but won't be opened")
    with pytest.raises(RuntimeError) as excinfo:
        run_workflow(
            artifacts_root=artifacts_root,
            input_path=str(docx),
            task="t",
            scaffold_version=1,
        )
    assert "workflow" in str(excinfo.value).lower()
