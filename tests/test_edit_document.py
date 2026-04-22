"""Tests for `scripts/edit_document.py` and helpers.

Fixtures are built programmatically in tmp_path so we never check binary
artifacts into the repo. Each test that needs a docx calls a `_make_*` helper
to materialise one with python-docx, then exercises the CLI subcommands as
either an importable function (preferred) or a subprocess (for end-to-end).
"""
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = REPO_ROOT / "scripts"

sys.path.insert(0, str(SCRIPTS_DIR))

import edit_document  # noqa: E402
from _docx_revisions import find_comment_runs, iter_comment_anchors  # noqa: E402
from _docx_selectors import parse_selector, resolve_selector  # noqa: E402


def _make_docx_two_comments(path: Path) -> None:
    from docx import Document

    doc = Document()
    p1 = doc.add_paragraph("The reviewer raised concerns about scalability.")
    doc.add_comment([p1.runs[0]], text="Cite a benchmark", author="Alice", initials="A")
    p2 = doc.add_paragraph("We will address this in section 3.")
    doc.add_comment([p2.runs[0]], text="Confirm scope first", author="Bob", initials="B")
    doc.save(str(path))


def _make_docx_no_comments(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Plain paragraph one.")
    doc.add_paragraph("Plain paragraph two.")
    doc.save(str(path))


def _make_docx_with_runs(path: Path) -> None:
    from docx import Document

    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Alpha ")
    para.add_run("Bravo ")
    para.add_run("Charlie")
    doc.save(str(path))


# ----- Selectors --------------------------------------------------------------


def test_parse_selector_paragraph():
    sel = parse_selector("paragraph:3")
    assert sel.kind == "paragraph"
    assert sel.paragraph_index == 3
    assert sel.run_index is None


def test_parse_selector_paragraph_run():
    sel = parse_selector("paragraph:2,run:1")
    assert sel.kind == "paragraph_run"
    assert sel.paragraph_index == 2
    assert sel.run_index == 1


def test_parse_selector_text_quoted():
    sel = parse_selector('text:"foo bar"')
    assert sel.kind == "text"
    assert sel.text_query == "foo bar"


def test_parse_selector_xpath():
    sel = parse_selector("xpath://w:p[1]")
    assert sel.kind == "xpath"
    assert sel.xpath_expr == "//w:p[1]"


def test_parse_selector_invalid():
    with pytest.raises(ValueError):
        parse_selector("nonsense:42")
    with pytest.raises(ValueError):
        parse_selector("paragraph:abc")


def test_resolve_selector_paragraph(tmp_path):
    path = tmp_path / "runs.docx"
    _make_docx_with_runs(path)
    from docx import Document

    doc = Document(str(path))
    para, run = resolve_selector(doc, parse_selector("paragraph:0"))
    assert para.text == "Alpha Bravo Charlie"
    assert run is None


def test_resolve_selector_paragraph_run(tmp_path):
    path = tmp_path / "runs.docx"
    _make_docx_with_runs(path)
    from docx import Document

    doc = Document(str(path))
    para, run = resolve_selector(doc, parse_selector("paragraph:0,run:1"))
    assert run is not None
    assert run.text.strip() == "Bravo"


def test_resolve_selector_text(tmp_path):
    path = tmp_path / "no_comments.docx"
    _make_docx_no_comments(path)
    from docx import Document

    doc = Document(str(path))
    para, _ = resolve_selector(doc, parse_selector('text:"two"'))
    assert "two" in para.text


# ----- read-comments ---------------------------------------------------------


def test_read_comments_returns_two_entries(tmp_path, capsys):
    path = tmp_path / "two.docx"
    _make_docx_two_comments(path)
    rc = edit_document.main(["read-comments", str(path), "--json"])
    assert rc == 0
    out = json.loads(capsys.readouterr().out)
    assert len(out) == 2
    assert {c["author"] for c in out} == {"Alice", "Bob"}
    bodies = {c["body"] for c in out}
    assert "Cite a benchmark" in bodies
    assert all(c["parent_id"] is None for c in out)
    # Anchor text matches the run we attached the comment to.
    alice = next(c for c in out if c["author"] == "Alice")
    assert "scalability" in alice["anchor_text"]
    assert alice["range"]["paragraph_indices"] == [0]


def test_no_comments_returns_empty(tmp_path, capsys):
    path = tmp_path / "empty.docx"
    _make_docx_no_comments(path)
    rc = edit_document.main(["read-comments", str(path), "--json"])
    assert rc == 0
    out = json.loads(capsys.readouterr().out)
    assert out == []


# ----- add-comment -----------------------------------------------------------


def test_add_comment_assigns_next_id_and_creates_part(tmp_path, capsys):
    path = tmp_path / "no.docx"
    _make_docx_no_comments(path)
    rc = edit_document.main(
        [
            "add-comment",
            str(path),
            "--selector",
            "paragraph:0",
            "--body",
            "fresh comment",
            "--author",
            "META",
            "--initials",
            "MC",
        ]
    )
    assert rc == 0
    capsys.readouterr()  # drain
    rc = edit_document.main(["read-comments", str(path), "--json"])
    out = json.loads(capsys.readouterr().out)
    assert len(out) == 1
    assert out[0]["author"] == "META"
    assert out[0]["body"] == "fresh comment"


# ----- reply-comment ---------------------------------------------------------


def test_reply_comment_inherits_parent_range(tmp_path, capsys):
    path = tmp_path / "two.docx"
    _make_docx_two_comments(path)
    capsys.readouterr()
    rc = edit_document.main(
        [
            "reply-comment",
            str(path),
            "--comment-id",
            "0",
            "--body",
            "Acknowledged",
            "--author",
            "META",
        ]
    )
    assert rc == 0
    reply_meta = json.loads(capsys.readouterr().out)
    assert reply_meta["parent_id"] == 0
    new_id = reply_meta["comment_id"]
    assert new_id != 0

    rc = edit_document.main(["read-comments", str(path), "--json"])
    out = json.loads(capsys.readouterr().out)
    reply = next(c for c in out if c["comment_id"] == new_id)
    assert reply["parent_id"] == 0
    assert reply["body"] == "Acknowledged"
    # Reply re-anchors to the same range as parent (paragraph 0).
    parent = next(c for c in out if c["comment_id"] == 0)
    assert reply["range"]["paragraph_indices"] == parent["range"]["paragraph_indices"]


def test_reply_comment_unknown_id_returns_error(tmp_path, capsys):
    path = tmp_path / "two.docx"
    _make_docx_two_comments(path)
    rc = edit_document.main(
        [
            "reply-comment",
            str(path),
            "--comment-id",
            "999",
            "--body",
            "x",
            "--author",
            "META",
        ]
    )
    assert rc == 2
    err = capsys.readouterr().err
    assert "999" in err


# ----- insert-tracked --------------------------------------------------------


def test_insert_tracked_emits_w_ins_with_author(tmp_path, capsys):
    path = tmp_path / "no.docx"
    _make_docx_no_comments(path)
    rc = edit_document.main(
        [
            "insert-tracked",
            str(path),
            "--selector",
            "paragraph:0",
            "--text",
            " (revised)",
            "--author",
            "META",
        ]
    )
    assert rc == 0
    capsys.readouterr()
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    body = doc.element.body
    ins_elements = list(body.iter(qn("w:ins")))
    assert len(ins_elements) == 1
    elem = ins_elements[0]
    assert elem.get(qn("w:author")) == "META"
    inserted_text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
    assert inserted_text == " (revised)"


# ----- delete-tracked --------------------------------------------------------


def test_delete_tracked_wraps_w_del_preserving_text(tmp_path, capsys):
    path = tmp_path / "runs.docx"
    _make_docx_with_runs(path)
    rc = edit_document.main(
        [
            "delete-tracked",
            str(path),
            "--selector",
            "paragraph:0,run:1",
            "--author",
            "META",
        ]
    )
    assert rc == 0
    capsys.readouterr()
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    body = doc.element.body
    dels = list(body.iter(qn("w:del")))
    assert len(dels) == 1
    deleted_text = "".join(t.text or "" for t in dels[0].iter(qn("w:delText")))
    assert deleted_text == "Bravo "


# ----- atomic save -----------------------------------------------------------


def test_atomic_save_preserves_path_and_no_temp_leftover(tmp_path, capsys):
    path = tmp_path / "no.docx"
    _make_docx_no_comments(path)
    pre_size = path.stat().st_size
    rc = edit_document.main(
        [
            "add-comment",
            str(path),
            "--selector",
            "paragraph:0",
            "--body",
            "hi",
            "--author",
            "X",
        ]
    )
    assert rc == 0
    capsys.readouterr()
    assert path.exists()
    assert path.stat().st_size > pre_size
    # No leftover temp files
    leftovers = list(tmp_path.glob("*.tmp.*"))
    assert leftovers == []


# ----- error handling --------------------------------------------------------


def test_corrupt_docx_returns_nonzero_with_clear_stderr(tmp_path, capsys):
    path = tmp_path / "bogus.docx"
    path.write_bytes(b"not a real docx")
    rc = edit_document.main(["read-comments", str(path), "--json"])
    assert rc == 1
    err = capsys.readouterr().err
    assert "ERROR:" in err


def test_missing_file_returns_nonzero(tmp_path, capsys):
    missing = tmp_path / "nope.docx"
    rc = edit_document.main(["read-comments", str(missing), "--json"])
    assert rc == 1
    err = capsys.readouterr().err
    assert "not found" in err


# ----- end-to-end via subprocess --------------------------------------------


def test_cli_subprocess_round_trip(tmp_path):
    path = tmp_path / "two.docx"
    _make_docx_two_comments(path)
    proc = subprocess.run(
        [sys.executable, str(SCRIPTS_DIR / "edit_document.py"), "read-comments", str(path), "--json"],
        capture_output=True,
        text=True,
        check=False,
    )
    assert proc.returncode == 0, proc.stderr
    payload = json.loads(proc.stdout)
    assert len(payload) == 2
