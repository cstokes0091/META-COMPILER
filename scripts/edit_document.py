#!/usr/bin/env python3
"""In-place .docx editor: comment read/reply, tracked insert/delete.

Subcommands
-----------
read-comments  <docx> [--json]
add-comment    <docx> --selector EXPR --body TEXT --author NAME [--initials XX]
reply-comment  <docx> --comment-id ID --body TEXT --author NAME [--initials XX]
insert-tracked <docx> --selector EXPR --text TEXT --author NAME
delete-tracked <docx> --selector EXPR --author NAME

Selector grammar
----------------
    paragraph:N            paragraph index N (0-based)
    paragraph:N,run:M      run M in paragraph N
    text:"..."             first paragraph containing the substring
    xpath:...              raw OOXML xpath returning <w:p>

Reply behaviour (v1)
--------------------
Word 365's threaded `commentsExtended.xml` part is out of scope. Replies are
written as a fresh comment anchored to the parent's range, with the body
prefixed by `[reply-to:<parent_id>] `. `read-comments` strips the marker and
populates `parent_id` on the JSON output. To accept/reject revisions, open the
document in Word; programmatic acceptance is not implemented in v1.

Atomic save
-----------
Each mutation writes to `<src>.tmp.<pid>`, fsyncs, then `os.replace()` over the
original. No lockfile - workflow conductors must serialize concurrent edits.
"""
from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path

from _docx_revisions import (
    delete_run_tracked,
    find_comment_runs,
    insert_tracked_text,
    iter_comment_anchors,
)
from _docx_selectors import parse_selector, resolve_selector

REPLY_MARKER_PREFIX = "[reply-to:"
REPLY_MARKER_SUFFIX = "] "
LARGE_FILE_WARN_BYTES = 25 * 1024 * 1024


def _strip_reply_marker(body: str) -> tuple[str, int | None]:
    if not body.startswith(REPLY_MARKER_PREFIX):
        return body, None
    close = body.find(REPLY_MARKER_SUFFIX)
    if close == -1:
        return body, None
    raw = body[len(REPLY_MARKER_PREFIX) : close]
    try:
        parent_id = int(raw)
    except ValueError:
        return body, None
    return body[close + len(REPLY_MARKER_SUFFIX) :], parent_id


def _atomic_save(doc, path: Path) -> None:
    tmp = path.with_suffix(path.suffix + f".tmp.{os.getpid()}")
    doc.save(str(tmp))
    fd = os.open(str(tmp), os.O_RDONLY)
    try:
        os.fsync(fd)
    finally:
        os.close(fd)
    os.replace(str(tmp), str(path))


def _load(path: Path):
    if not path.exists():
        raise FileNotFoundError(f"file not found: {path}")
    if path.stat().st_size > LARGE_FILE_WARN_BYTES:
        print(
            f"WARNING: {path} is larger than {LARGE_FILE_WARN_BYTES // (1024 * 1024)} MB; "
            "performance not optimised in v1.",
            file=sys.stderr,
        )
    from docx import Document

    return Document(str(path))


def cmd_read_comments(args: argparse.Namespace) -> int:
    doc = _load(Path(args.docx))
    paragraphs = list(doc.paragraphs)
    anchor_paras = {cid: paras for cid, paras in iter_comment_anchors(doc)}

    out: list[dict] = []
    for comment in doc.comments:
        runs = find_comment_runs(doc, comment.comment_id)
        anchor_text = "".join(r.text for r in runs)
        body_text = comment.text or ""
        body_clean, parent_id = _strip_reply_marker(body_text)
        para_idxs = anchor_paras.get(comment.comment_id, [])
        first_para = paragraphs[para_idxs[0]] if para_idxs else None
        if first_para is not None and anchor_text:
            try:
                start = first_para.text.index(anchor_text)
                end = start + len(anchor_text)
            except ValueError:
                start, end = -1, -1
        else:
            start, end = -1, -1
        out.append(
            {
                "comment_id": comment.comment_id,
                "anchor_text": anchor_text,
                "range": {
                    "paragraph_indices": para_idxs,
                    "start": start,
                    "end": end,
                },
                "author": comment.author,
                "initials": comment.initials,
                "timestamp_iso": comment.timestamp.isoformat() if comment.timestamp else None,
                "body": body_clean,
                "parent_id": parent_id,
            }
        )
    if args.json or not sys.stdout.isatty():
        print(json.dumps(out, indent=2, default=str))
    else:
        for c in out:
            parent = f" (reply to #{c['parent_id']})" if c["parent_id"] is not None else ""
            print(f"#{c['comment_id']} {c['author']!r}{parent}: {c['body']}")
            print(f"   anchor: {c['anchor_text']!r}  paras={c['range']['paragraph_indices']}")
    return 0


def cmd_add_comment(args: argparse.Namespace) -> int:
    path = Path(args.docx)
    doc = _load(path)
    selector = parse_selector(args.selector)
    para, run = resolve_selector(doc, selector)
    targets = [run] if run is not None else list(para.runs)
    if not targets:
        run = para.add_run("")
        targets = [run]
    doc.add_comment(
        runs=targets,
        text=args.body,
        author=args.author,
        initials=args.initials or "",
    )
    _atomic_save(doc, path)
    new_id = max((c.comment_id for c in doc.comments), default=-1)
    print(json.dumps({"status": "ok", "comment_id": new_id}))
    return 0


def cmd_reply_comment(args: argparse.Namespace) -> int:
    path = Path(args.docx)
    doc = _load(path)
    parent = doc.comments.get(args.comment_id)
    if parent is None:
        print(f"ERROR: parent comment id {args.comment_id} not found", file=sys.stderr)
        return 2
    runs = find_comment_runs(doc, args.comment_id)
    if not runs:
        print(
            f"ERROR: cannot locate anchor range for comment id {args.comment_id}",
            file=sys.stderr,
        )
        return 2
    body = f"{REPLY_MARKER_PREFIX}{args.comment_id}{REPLY_MARKER_SUFFIX}{args.body}"
    doc.add_comment(
        runs=runs,
        text=body,
        author=args.author,
        initials=args.initials or "",
    )
    _atomic_save(doc, path)
    new_id = max((c.comment_id for c in doc.comments), default=-1)
    print(json.dumps({"status": "ok", "comment_id": new_id, "parent_id": args.comment_id}))
    return 0


def cmd_insert_tracked(args: argparse.Namespace) -> int:
    path = Path(args.docx)
    doc = _load(path)
    selector = parse_selector(args.selector)
    para, _run = resolve_selector(doc, selector)
    insert_tracked_text(para, args.text, args.author)
    _atomic_save(doc, path)
    print(json.dumps({"status": "ok"}))
    return 0


def cmd_delete_tracked(args: argparse.Namespace) -> int:
    path = Path(args.docx)
    doc = _load(path)
    selector = parse_selector(args.selector)
    para, run = resolve_selector(doc, selector)
    if run is None:
        runs = list(para.runs)
        if not runs:
            print("ERROR: paragraph has no runs to delete", file=sys.stderr)
            return 2
        for r in runs:
            delete_run_tracked(r, args.author)
    else:
        delete_run_tracked(run, args.author)
    _atomic_save(doc, path)
    print(json.dumps({"status": "ok"}))
    return 0


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="edit_document",
        description="In-place .docx editor for comments and tracked changes.",
        epilog=(
            "Out of scope for v1: accept-revision, reject-revision, threaded comment "
            "replies via commentsExtended.xml. Replies use a flat anchor with a "
            "[reply-to:<id>] marker; round-trips through read-comments."
        ),
    )
    sub = parser.add_subparsers(dest="command", required=True)

    rc = sub.add_parser("read-comments", help="Emit comments as JSON or human-readable text.")
    rc.add_argument("docx")
    rc.add_argument("--json", action="store_true", help="Force JSON output even on a TTY.")
    rc.set_defaults(func=cmd_read_comments)

    ac = sub.add_parser("add-comment", help="Add a new comment anchored to a selector.")
    ac.add_argument("docx")
    ac.add_argument("--selector", required=True)
    ac.add_argument("--body", required=True)
    ac.add_argument("--author", required=True)
    ac.add_argument("--initials", default=None)
    ac.set_defaults(func=cmd_add_comment)

    rp = sub.add_parser("reply-comment", help="Reply (flat) to an existing comment.")
    rp.add_argument("docx")
    rp.add_argument("--comment-id", required=True, type=int)
    rp.add_argument("--body", required=True)
    rp.add_argument("--author", required=True)
    rp.add_argument("--initials", default=None)
    rp.set_defaults(func=cmd_reply_comment)

    ins = sub.add_parser("insert-tracked", help="Append tracked-insert text to the selected paragraph.")
    ins.add_argument("docx")
    ins.add_argument("--selector", required=True)
    ins.add_argument("--text", required=True)
    ins.add_argument("--author", required=True)
    ins.set_defaults(func=cmd_insert_tracked)

    dele = sub.add_parser("delete-tracked", help="Wrap selected run(s) in a tracked-delete revision.")
    dele.add_argument("docx")
    dele.add_argument("--selector", required=True)
    dele.add_argument("--author", required=True)
    dele.set_defaults(func=cmd_delete_tracked)

    return parser


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)
    try:
        return args.func(args)
    except (FileNotFoundError, LookupError, IndexError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: unexpected: {exc.__class__.__name__}: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
