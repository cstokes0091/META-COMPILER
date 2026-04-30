"""Stage 4 final-synthesis sub-stage.

After the @execution-orchestrator ralph loop populates
`executions/v{N}/work/<capability_id>/`, the per-capability output is a
collection of fragments — not a deliverable. This sub-stage assembles those
fragments into a unified, project-type-aware artifact:

* `algorithm`/`hybrid` → `final/library/<package>/...` (importable Python package)
* `report`/`hybrid`    → `final/document/<slug>.md` + .docx (single document)
* `workflow`           → `final/application/run.py` + bucket layout (runnable app)

The split mirrors the wiki-reconcile / cross-source-synthesis pattern: a
preflight CLI (`final-synthesize-start`) walks `work/` and emits a
deterministic work plan, an orchestrator agent fans out per-modality
synthesizer subagents whose JSON returns are persisted verbatim, and a
postflight CLI (`final-synthesize-finalize`) validates each return and
materializes `executions/v{N}/final/<bucket>/` deterministically.

The synthesizer subagents never write files. The CLI never picks names,
decides layouts, or composes prose.
"""
from __future__ import annotations

import json
import re
import shutil
from pathlib import Path
from typing import Any

from ..artifacts import (
    ArtifactPaths,
    build_paths,
    ensure_layout,
    latest_decision_log_path,
    latest_scaffold_path,
    load_manifest,
    save_manifest,
)
from ..io import dump_yaml, load_yaml
from ..project_types import scaffold_subdirs_for
from ..utils import iso_now, slugify
from ..validation import (
    validate_application_synthesis_return,
    validate_document_synthesis_return,
    validate_library_synthesis_return,
)


# ---------------------------------------------------------------------------
# Module constants
# ---------------------------------------------------------------------------


_CODE_SUFFIXES: frozenset[str] = frozenset({
    ".py", ".js", ".ts", ".tsx", ".jsx", ".rs", ".go", ".java", ".kt",
    ".c", ".cc", ".cpp", ".h", ".hpp", ".rb", ".php", ".cs", ".swift",
    ".sh", ".bash", ".zsh", ".sql", ".m", ".mm", ".scala", ".jl",
})
_DOC_SUFFIXES: frozenset[str] = frozenset({".md", ".rst", ".txt", ".tex"})
_DATA_SUFFIXES: frozenset[str] = frozenset({".yaml", ".yml", ".json", ".csv", ".tsv"})
# Files written by the orchestrator/planner/reviewer themselves — never the
# capability's actual deliverable. Excluded from synthesis input.
_CAPABILITY_BUCKET_FILES: frozenset[str] = frozenset({
    "_plan.yaml", "_verdict.yaml", "_manifest.yaml",
})

# project_type → (modality_keys, ...). Branches the synthesis fan-out.
_PROJECT_TYPE_MODALITIES: dict[str, tuple[str, ...]] = {
    "algorithm": ("library",),
    "report": ("document",),
    "hybrid": ("library", "document"),
    "workflow": ("application",),
}

# Per-modality fragment filter. The library synthesizer only sees code; the
# document synthesizer only sees prose; the application synthesizer sees
# everything (it's wiring fragments into a runnable layout).
_MODALITY_FRAGMENT_FILTER: dict[str, frozenset[str]] = {
    "library": frozenset({"code"}),
    "document": frozenset({"document"}),
    "application": frozenset({"code", "document", "data"}),
}

_REQ_RE = re.compile(r"REQ-\d{3}")


# ---------------------------------------------------------------------------
# Fragment classification
# ---------------------------------------------------------------------------


def _classify_modality(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in _CODE_SUFFIXES:
        return "code"
    if suffix in _DOC_SUFFIXES:
        return "document"
    if suffix in _DATA_SUFFIXES:
        return "data"
    return "other"


def _read_fragment_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except (OSError, UnicodeDecodeError):
        return ""


def _walk_work_fragments(work_dir: Path, root: Path) -> list[dict[str, Any]]:
    """Walk every per-capability subdir of work_dir and classify each file.

    Capability-bucket files (_plan.yaml, _verdict.yaml, _manifest.yaml) are
    skipped — they're orchestrator/reviewer scaffolding, not deliverables.
    """
    fragments: list[dict[str, Any]] = []
    if not work_dir.exists():
        return fragments
    for capability_dir in sorted(work_dir.iterdir()):
        if not capability_dir.is_dir():
            continue
        capability_id = capability_dir.name
        for path in sorted(capability_dir.rglob("*")):
            if not path.is_file():
                continue
            if path.name in _CAPABILITY_BUCKET_FILES:
                continue
            modality = _classify_modality(path)
            relative = path.relative_to(capability_dir)
            text = _read_fragment_text(path)
            req_mentions = sorted(set(_REQ_RE.findall(text)))
            try:
                line_count = sum(1 for _ in text.splitlines())
            except Exception:
                line_count = 0
            try:
                size_bytes = path.stat().st_size
            except OSError:
                size_bytes = 0
            try:
                rel_to_root = path.relative_to(root)
                root_relative = str(rel_to_root)
            except ValueError:
                root_relative = str(path)
            fragments.append(
                {
                    "capability": capability_id,
                    "relative_path": str(relative),
                    "absolute_path": str(path),
                    "root_relative_path": root_relative,
                    "modality": modality,
                    "size_bytes": size_bytes,
                    "line_count": line_count,
                    "req_mentions": req_mentions,
                }
            )
    return fragments


def _modalities_for_project_type(project_type: str) -> tuple[str, ...]:
    return _PROJECT_TYPE_MODALITIES.get(project_type, ())


def _slice_fragments_for_modality(
    modality: str, fragments: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    accept = _MODALITY_FRAGMENT_FILTER.get(modality, frozenset())
    return [f for f in fragments if f["modality"] in accept]


def _fragment_token(fragment: dict[str, Any]) -> str:
    return f"{fragment['capability']}:{fragment['relative_path']}"


def _expected_citation_ids(paths: ArtifactPaths) -> set[str]:
    payload = load_yaml(paths.citations_index_path) if paths.citations_index_path.exists() else {}
    if not isinstance(payload, dict):
        return set()
    citations = payload.get("citations") or {}
    if not isinstance(citations, dict):
        return set()
    return {str(k) for k in citations.keys() if isinstance(k, str) and k.strip()}


def _expected_req_ids(scaffold_root: Path, decision_log: dict[str, Any]) -> set[str]:
    """Pull REQ-NNN ids from REQ_TRACE.yaml (preferred) or the decision log."""
    trace_path = scaffold_root / "verification" / "REQ_TRACE.yaml"
    if trace_path.exists():
        payload = load_yaml(trace_path) or {}
        if isinstance(payload, dict):
            trace = payload.get("req_trace") or payload
            entries = trace.get("entries") if isinstance(trace, dict) else None
            if isinstance(entries, list):
                ids = {
                    str(row.get("req_id"))
                    for row in entries
                    if isinstance(row, dict)
                    and isinstance(row.get("req_id"), str)
                    and _REQ_RE.fullmatch(str(row.get("req_id")))
                }
                if ids:
                    return ids
    # Fallback: scan the decision log
    root = decision_log.get("decision_log") or {}
    requirements = root.get("requirements") or []
    return {
        str(req.get("id"))
        for req in requirements
        if isinstance(req, dict)
        and isinstance(req.get("id"), str)
        and _REQ_RE.fullmatch(str(req.get("id")))
    }


def _expected_workflow_buckets(project_type: str) -> set[str]:
    """For the application modality. The synthesizer's `directory_layout`
    must include every bucket the project_type promises. We add `orchestrator`
    as a supporting subdirectory the synthesizer is expected to populate."""
    if project_type != "workflow":
        return set()
    buckets = set(scaffold_subdirs_for("workflow"))
    buckets.add("orchestrator")
    return buckets


# ---------------------------------------------------------------------------
# Shared bootstrap helpers
# ---------------------------------------------------------------------------


def _resolve_decision_log_version(paths: ArtifactPaths, requested: int | None) -> int:
    if requested is not None:
        return requested
    latest = latest_decision_log_path(paths)
    if latest is None:
        raise RuntimeError("No decision log found. Run elicit-vision first.")
    return latest[0]


def _resolve_scaffold(paths: ArtifactPaths, version: int) -> Path:
    candidate = paths.scaffolds_dir / f"v{version}"
    if candidate.exists():
        return candidate
    latest = latest_scaffold_path(paths)
    if latest is None:
        raise RuntimeError("No scaffold found. Run scaffold first.")
    if latest[0] != version:
        raise RuntimeError(
            f"Latest scaffold version (v{latest[0]}) does not match Decision Log v{version}."
        )
    return latest[1]


def _load_decision_log(paths: ArtifactPaths, version: int) -> dict[str, Any]:
    decision_log_path = paths.decision_logs_dir / f"decision_log_v{version}.yaml"
    if not decision_log_path.exists():
        raise RuntimeError(f"Decision log missing: {decision_log_path}")
    return load_yaml(decision_log_path) or {}


def _project_type_from_execution_manifest(scaffold_root: Path) -> str:
    execution_manifest_path = scaffold_root / "EXECUTION_MANIFEST.yaml"
    if not execution_manifest_path.exists():
        raise RuntimeError(f"Execution manifest missing: {execution_manifest_path}")
    payload = load_yaml(execution_manifest_path) or {}
    execution_root = payload.get("execution") or {}
    project_type = str(execution_root.get("project_type") or "")
    if not project_type:
        raise RuntimeError(
            f"EXECUTION_MANIFEST.yaml at {execution_manifest_path} missing project_type"
        )
    return project_type


def _project_slug(decision_log: dict[str, Any]) -> str:
    root = decision_log.get("decision_log") or {}
    meta = root.get("meta") or {}
    name = meta.get("project_name") or "synthesis"
    return slugify(str(name)) or "synthesis"


def _set_manifest_stage(paths: ArtifactPaths, last_completed_stage: str) -> None:
    manifest = load_manifest(paths)
    if not manifest:
        return
    wm = manifest.setdefault("workspace_manifest", {})
    research = wm.setdefault("research", {})
    research["last_completed_stage"] = last_completed_stage
    save_manifest(paths, manifest)


# ---------------------------------------------------------------------------
# Preflight: final-synthesize-start
# ---------------------------------------------------------------------------


def run_final_synthesize_start(
    artifacts_root: Path,
    workspace_root: Path,
    decision_log_version: int | None = None,
) -> dict[str, Any]:
    """Walk executions/v{N}/work/, classify fragments, write the work plan
    + synthesis_request that the @final-synthesis-orchestrator consumes."""
    paths = build_paths(artifacts_root)
    ensure_layout(paths)

    version = _resolve_decision_log_version(paths, decision_log_version)
    scaffold_root = _resolve_scaffold(paths, version)
    project_type = _project_type_from_execution_manifest(scaffold_root)
    decision_log = _load_decision_log(paths, version)

    output_dir = paths.executions_dir / f"v{version}"
    work_dir = output_dir / "work"
    final_dir = paths.final_dir_for(version)

    if not work_dir.exists() or not any(work_dir.rglob("*")):
        raise RuntimeError(
            f"executions/v{version}/work/ is empty. Run "
            "`meta-compiler phase4-finalize --start` and let "
            "@execution-orchestrator populate the work directory before "
            "running final-synthesize-start."
        )

    fragments = _walk_work_fragments(work_dir, paths.root)
    modality_keys = _modalities_for_project_type(project_type)
    if not modality_keys:
        raise RuntimeError(
            f"Unknown project_type {project_type!r}; cannot determine synthesis modalities."
        )

    expected_citation_ids = sorted(_expected_citation_ids(paths))
    expected_req_ids = sorted(_expected_req_ids(scaffold_root, decision_log))
    workflow_buckets = sorted(_expected_workflow_buckets(project_type))
    project_slug = _project_slug(decision_log)

    modalities_payload: dict[str, Any] = {}
    for modality in modality_keys:
        sliced = _slice_fragments_for_modality(modality, fragments)
        tokens = sorted({_fragment_token(f) for f in sliced})
        modality_output_dir = {
            "library": final_dir / "library",
            "document": final_dir / "document",
            "application": final_dir / "application",
        }[modality]
        modalities_payload[modality] = {
            "fragments": sliced,
            "expected_fragment_tokens": tokens,
            "fragment_count": len(sliced),
            "output_dir": str(modality_output_dir),
            "subagent_return_path": str(
                paths.final_synthesis_subagent_returns_dir / f"{modality}.json"
            ),
        }

    generated_at = iso_now()

    # Change D: thread CONTEXT.md (Stage 3 output) into the synthesis
    # request so per-modality synthesizers honor the same Domain +
    # Architecture vocabulary as the work-loop palette agents. CONTEXT.md
    # may not exist on legacy scaffolds; surface as None and let the
    # orchestrator handle gracefully.
    context_md_path = scaffold_root / "CONTEXT.md"
    context_md_relative = (
        str(context_md_path.relative_to(paths.root))
        if context_md_path.exists()
        else None
    )

    plan_payload = {
        "final_synthesis_work_plan": {
            "generated_at": generated_at,
            "decision_log_version": version,
            "project_type": project_type,
            "project_slug": project_slug,
            "scaffold_root": str(scaffold_root),
            "work_dir": str(work_dir),
            "final_dir": str(final_dir),
            "context_md_path": context_md_relative,
            "modality_keys": list(modality_keys),
            "modalities": modalities_payload,
            "expected_citation_ids": expected_citation_ids,
            "expected_req_ids": expected_req_ids,
            "workflow_buckets": workflow_buckets,
            "fragment_count": len(fragments),
        }
    }
    dump_yaml(paths.final_synthesis_work_plan_path, plan_payload)

    request_payload = {
        "final_synthesis_request": {
            "generated_at": generated_at,
            "decision_log_version": version,
            "project_type": project_type,
            "work_plan_path": str(paths.final_synthesis_work_plan_path),
            "subagent_returns_dir": str(paths.final_synthesis_subagent_returns_dir),
            "context_md_path": context_md_relative,
            "modality_keys": list(modality_keys),
            "next_action": (
                "Invoke .github/prompts/final-synthesis.prompt.md to fan out "
                "per-modality synthesizer subagents, persist each return to "
                "subagent_returns/<modality>.json, then run "
                "`meta-compiler final-synthesize-finalize`. Each synthesizer "
                "MUST read context_md_path first (when set) to honor the "
                "same Domain + Architecture vocabulary as the work-loop "
                "palette agents."
            ),
        }
    }
    dump_yaml(paths.final_synthesis_request_path, request_payload)

    _set_manifest_stage(paths, "4-synthesis-pending")

    return {
        "status": "ready_for_orchestrator",
        "decision_log_version": version,
        "project_type": project_type,
        "modality_keys": list(modality_keys),
        "fragment_count": len(fragments),
        "fragments_per_modality": {
            mod: modalities_payload[mod]["fragment_count"] for mod in modality_keys
        },
        "work_plan_path": str(paths.final_synthesis_work_plan_path),
        "request_path": str(paths.final_synthesis_request_path),
        "subagent_returns_dir": str(paths.final_synthesis_subagent_returns_dir),
        "final_dir": str(final_dir),
    }


# ---------------------------------------------------------------------------
# Postflight: final-synthesize-finalize
# ---------------------------------------------------------------------------


def _load_subagent_return(paths: ArtifactPaths, modality: str) -> dict[str, Any] | None:
    path = paths.final_synthesis_subagent_returns_dir / f"{modality}.json"
    if not path.exists():
        return None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(
            f"subagent_returns/{modality}.json is unreadable: {exc}"
        ) from exc
    if not isinstance(payload, dict):
        raise ValueError(
            f"subagent_returns/{modality}.json: must be a JSON object"
        )
    return payload


def _validate_returns(
    work_plan: dict[str, Any],
    returns: dict[str, dict[str, Any]],
) -> dict[str, list[str]]:
    """Validate every return; return {modality: issues[]}."""
    issues_per_modality: dict[str, list[str]] = {}

    expected_citation_ids = set(work_plan.get("expected_citation_ids") or [])
    expected_req_ids = set(work_plan.get("expected_req_ids") or [])
    workflow_buckets = set(work_plan.get("workflow_buckets") or [])

    for modality, payload in returns.items():
        slice_payload = (work_plan.get("modalities") or {}).get(modality, {})
        expected_fragments = set(slice_payload.get("expected_fragment_tokens") or [])
        if modality == "library":
            issues_per_modality[modality] = validate_library_synthesis_return(
                payload,
                expected_fragments=expected_fragments,
                expected_req_ids=expected_req_ids,
            )
        elif modality == "document":
            issues_per_modality[modality] = validate_document_synthesis_return(
                payload,
                expected_fragments=expected_fragments,
                expected_citation_ids=expected_citation_ids,
                expected_req_ids=expected_req_ids,
            )
        elif modality == "application":
            issues_per_modality[modality] = validate_application_synthesis_return(
                payload,
                expected_fragments=expected_fragments,
                expected_buckets=workflow_buckets,
                expected_req_ids=expected_req_ids,
            )
        else:
            issues_per_modality[modality] = [
                f"unknown modality {modality!r} in subagent_returns"
            ]
    return issues_per_modality


def _ensure_clean_tmp(tmp_dir: Path) -> None:
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir)
    tmp_dir.mkdir(parents=True, exist_ok=True)


def _swap_tmp_into_place(tmp_dir: Path, final_dir: Path) -> None:
    """Atomically replace `final_dir` with `tmp_dir` contents.

    `final_dir` may not exist yet. If it does, it is removed first. The
    swap is one rename — interrupted runs leave the previous `final_dir`
    intact and `tmp_dir` orphaned for the next start to clean up.
    """
    if final_dir.exists():
        shutil.rmtree(final_dir)
    final_dir.parent.mkdir(parents=True, exist_ok=True)
    tmp_dir.rename(final_dir)


def _fragment_lookup(work_dir: Path, fragments: list[dict[str, Any]]) -> dict[str, Path]:
    """Build a `<capability>:<relative_path>` → absolute Path lookup."""
    out: dict[str, Path] = {}
    for fragment in fragments:
        token = _fragment_token(fragment)
        absolute = Path(fragment.get("absolute_path") or "")
        if not absolute.is_absolute():
            absolute = work_dir / fragment["capability"] / fragment["relative_path"]
        out[token] = absolute
    return out


# ---------------------------------------------------------------------------
# Apply: library
# ---------------------------------------------------------------------------


def _apply_library_proposal(
    payload: dict[str, Any],
    work_plan_slice: dict[str, Any],
    work_dir: Path,
    final_tmp_dir: Path,
) -> dict[str, Any]:
    library_root = final_tmp_dir / "library"
    package_name = str(payload["package_name"])
    package_root = library_root / package_name
    package_root.mkdir(parents=True, exist_ok=True)

    fragment_lookup = _fragment_lookup(work_dir, work_plan_slice.get("fragments") or [])

    written_files: list[str] = []

    for entry in payload.get("module_layout") or []:
        target = library_root / entry["target_path"]
        target.parent.mkdir(parents=True, exist_ok=True)
        chunks: list[str] = []
        header = entry.get("header_prose")
        if isinstance(header, str) and header.strip():
            chunks.append(header.rstrip() + "\n")
        for source in entry.get("sources") or []:
            token = _fragment_token(source)
            fragment_path = fragment_lookup.get(token)
            if fragment_path is None or not fragment_path.exists():
                raise FileNotFoundError(
                    f"library_synthesis: fragment {token!r} resolved to "
                    f"{fragment_path} but the file is missing"
                )
            chunks.append(fragment_path.read_text(encoding="utf-8").rstrip() + "\n")
        footer = entry.get("footer_prose")
        if isinstance(footer, str) and footer.strip():
            chunks.append(footer.rstrip() + "\n")
        target.write_text("\n".join(chunks).rstrip() + "\n", encoding="utf-8")
        written_files.append(str(target.relative_to(final_tmp_dir.parent)))

    # __init__.py from exports + a re-export shim. The synthesizer may opt
    # out by listing a "<package>/__init__.py" target_path explicitly; in
    # that case we skip the auto-init.
    init_target = package_root / "__init__.py"
    if not init_target.exists():
        exports = [
            str(s) for s in (payload.get("exports") or [])
            if isinstance(s, str) and s.strip()
        ]
        init_lines = [f'"""{package_name} — synthesized package."""', ""]
        if exports:
            init_lines.append("__all__ = [")
            for sym in exports:
                init_lines.append(f"    {sym!r},")
            init_lines.append("]")
        init_target.write_text("\n".join(init_lines) + "\n", encoding="utf-8")
        written_files.append(str(init_target.relative_to(final_tmp_dir.parent)))

    # README.md
    readme_target = library_root / "README.md"
    readme_lines: list[str] = [f"# {package_name}", ""]
    for section in payload.get("readme_sections") or []:
        heading = str(section.get("heading") or "").strip()
        body = str(section.get("body") or "").strip()
        if not heading or not body:
            continue
        readme_lines.append(f"## {heading}")
        readme_lines.append("")
        readme_lines.append(body)
        readme_lines.append("")
    readme_target.write_text("\n".join(readme_lines).rstrip() + "\n", encoding="utf-8")
    written_files.append(str(readme_target.relative_to(final_tmp_dir.parent)))

    # Optional pyproject.toml
    metadata = payload.get("package_metadata")
    if isinstance(metadata, dict):
        dist_name = str(metadata.get("name") or package_name).strip()
        description = str(metadata.get("description") or f"{package_name} synthesized package").strip()
        python_requires = str(metadata.get("python_requires") or ">=3.10").strip()
        pyproject_lines = [
            "[build-system]",
            'requires = ["setuptools>=68"]',
            'build-backend = "setuptools.build_meta"',
            "",
            "[project]",
            f'name = "{dist_name}"',
            'version = "0.1.0"',
            f'description = "{description}"',
            f'requires-python = "{python_requires}"',
        ]
        entry_points = payload.get("entry_points") or []
        if isinstance(entry_points, list) and entry_points:
            pyproject_lines.extend(["", "[project.scripts]"])
            for entry in entry_points:
                if not isinstance(entry, dict):
                    continue
                name = str(entry.get("name") or "").strip()
                target = str(entry.get("target") or "").strip()
                if name and target:
                    pyproject_lines.append(f'{name} = "{target}"')
        pyproject_target = library_root / "pyproject.toml"
        pyproject_target.write_text("\n".join(pyproject_lines) + "\n", encoding="utf-8")
        written_files.append(str(pyproject_target.relative_to(final_tmp_dir.parent)))

    return {
        "modality": "library",
        "package_name": package_name,
        "files_written": written_files,
        "module_count": len(payload.get("module_layout") or []),
        "exports_count": len(payload.get("exports") or []),
    }


# ---------------------------------------------------------------------------
# Apply: document
# ---------------------------------------------------------------------------


def _render_document_markdown(
    payload: dict[str, Any],
    work_plan_slice: dict[str, Any],
    work_dir: Path,
) -> tuple[str, str]:
    """Build the `<title>.md` body and a separate `references.md` body."""
    fragment_lookup = _fragment_lookup(work_dir, work_plan_slice.get("fragments") or [])

    title = str(payload.get("title") or "Synthesized Document").strip()
    abstract = str(payload.get("abstract") or "").strip()
    intro = str(payload.get("intro_prose") or "").strip()
    conclusion = str(payload.get("conclusion_prose") or "").strip()

    lines: list[str] = [f"# {title}", ""]
    if abstract:
        lines.append(f"_{abstract}_")
        lines.append("")
    if intro:
        lines.append(intro)
        lines.append("")

    for section in payload.get("section_order") or []:
        if not isinstance(section, dict):
            continue
        heading = str(section.get("heading") or "").strip()
        if not heading:
            continue
        lines.append(f"## {heading}")
        lines.append("")
        source = section.get("source") or {}
        synth_prose = source.get("synthesizer_prose")
        capability = source.get("capability")
        rel = source.get("file") or source.get("relative_path") or source.get("path")
        if isinstance(capability, str) and isinstance(rel, str) and capability.strip() and rel.strip():
            token = f"{capability}:{rel}"
            fragment_path = fragment_lookup.get(token)
            body_text = ""
            if fragment_path is not None and fragment_path.exists():
                body_text = fragment_path.read_text(encoding="utf-8").strip()
                # Strip a leading H1 if present so we don't double-header.
                body_lines = body_text.splitlines()
                if body_lines and body_lines[0].startswith("# "):
                    body_text = "\n".join(body_lines[1:]).lstrip()
            lines.append(body_text)
        elif isinstance(synth_prose, str) and synth_prose.strip():
            lines.append(synth_prose.strip())
        lines.append("")
        transitions = section.get("transitions_after")
        if isinstance(transitions, str) and transitions.strip():
            lines.append(transitions.strip())
            lines.append("")

    if conclusion:
        lines.append("## Conclusion")
        lines.append("")
        lines.append(conclusion)
        lines.append("")

    references_unified = payload.get("references_unified") or []
    if references_unified:
        lines.append("## References")
        lines.append("")
        for entry in references_unified:
            if not isinstance(entry, dict):
                continue
            cid = str(entry.get("id") or "").strip()
            human = str(entry.get("human") or "").strip()
            if not cid or not human:
                continue
            lines.append(f"- **{cid}** — {human}")
        lines.append("")

    document_md = "\n".join(lines).rstrip() + "\n"

    references_lines: list[str] = ["# References", ""]
    for entry in references_unified:
        if not isinstance(entry, dict):
            continue
        cid = str(entry.get("id") or "").strip()
        human = str(entry.get("human") or "").strip()
        if not cid or not human:
            continue
        references_lines.append(f"- **{cid}** — {human}")
    references_md = "\n".join(references_lines).rstrip() + "\n"

    return document_md, references_md


def _try_render_docx(markdown_path: Path, docx_path: Path, title: str) -> bool:
    """Attempt to render <document>.md → <document>.docx via scripts/write_document.py.

    Returns True on success. Returns False (with a soft warning) when the
    `docx` package is unavailable — the markdown is still the canonical
    artifact, the docx is a convenience.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return False
    text = markdown_path.read_text(encoding="utf-8")
    doc = Document()
    doc.add_heading(title, level=0)
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return True


def _apply_document_proposal(
    payload: dict[str, Any],
    work_plan_slice: dict[str, Any],
    work_dir: Path,
    final_tmp_dir: Path,
    project_slug: str,
) -> dict[str, Any]:
    document_root = final_tmp_dir / "document"
    document_root.mkdir(parents=True, exist_ok=True)

    document_md, references_md = _render_document_markdown(
        payload, work_plan_slice, work_dir
    )

    md_target = document_root / f"{project_slug}.md"
    md_target.write_text(document_md, encoding="utf-8")
    written_files = [str(md_target.relative_to(final_tmp_dir.parent))]

    references_target = document_root / "references.md"
    references_target.write_text(references_md, encoding="utf-8")
    written_files.append(str(references_target.relative_to(final_tmp_dir.parent)))

    docx_target = document_root / f"{project_slug}.docx"
    title = str(payload.get("title") or project_slug)
    docx_rendered = _try_render_docx(md_target, docx_target, title)
    if docx_rendered:
        written_files.append(str(docx_target.relative_to(final_tmp_dir.parent)))

    return {
        "modality": "document",
        "title": title,
        "files_written": written_files,
        "section_count": len(payload.get("section_order") or []),
        "docx_rendered": docx_rendered,
    }


# ---------------------------------------------------------------------------
# Apply: application
# ---------------------------------------------------------------------------


def _apply_application_proposal(
    payload: dict[str, Any],
    work_plan_slice: dict[str, Any],
    work_dir: Path,
    final_tmp_dir: Path,
) -> dict[str, Any]:
    application_root = final_tmp_dir / "application"
    application_root.mkdir(parents=True, exist_ok=True)

    fragment_lookup = _fragment_lookup(work_dir, work_plan_slice.get("fragments") or [])

    written_files: list[str] = []

    directory_layout = payload.get("directory_layout") or {}
    for bucket, entries in directory_layout.items():
        if not isinstance(entries, list):
            continue
        bucket_root = application_root / str(bucket)
        bucket_root.mkdir(parents=True, exist_ok=True)
        for entry in entries:
            if not isinstance(entry, dict):
                continue
            source = str(entry.get("source") or "")
            target = str(entry.get("target") or "").strip()
            if not source or not target:
                continue
            fragment_path = fragment_lookup.get(source)
            if fragment_path is None or not fragment_path.exists():
                raise FileNotFoundError(
                    f"application_synthesis: fragment {source!r} resolved to "
                    f"{fragment_path} but the file is missing"
                )
            destination = application_root / target
            destination.parent.mkdir(parents=True, exist_ok=True)
            destination.write_bytes(fragment_path.read_bytes())
            written_files.append(str(destination.relative_to(final_tmp_dir.parent)))

    entry_point = payload.get("entry_point") or {}
    filename = str(entry_point.get("filename") or "run.py")
    body = str(entry_point.get("body") or "").rstrip() + "\n"
    entry_target = application_root / filename
    entry_target.write_text(body, encoding="utf-8")
    written_files.append(str(entry_target.relative_to(final_tmp_dir.parent)))

    # README.md
    readme_target = application_root / "README.md"
    app_name = str(payload.get("application_name") or "synthesized-application")
    readme_lines: list[str] = [f"# {app_name}", ""]
    invocation = str(entry_point.get("invocation") or "").strip()
    if invocation:
        readme_lines.extend([
            "## Run",
            "",
            "```",
            invocation,
            "```",
            "",
        ])
    for section in payload.get("readme_sections") or []:
        heading = str(section.get("heading") or "").strip()
        body_text = str(section.get("body") or "").strip()
        if not heading or not body_text:
            continue
        if heading == "Run" and invocation:
            continue  # already inserted above
        readme_lines.append(f"## {heading}")
        readme_lines.append("")
        readme_lines.append(body_text)
        readme_lines.append("")
    readme_target.write_text("\n".join(readme_lines).rstrip() + "\n", encoding="utf-8")
    written_files.append(str(readme_target.relative_to(final_tmp_dir.parent)))

    # requirements.txt
    dependencies = [
        str(d).strip() for d in (payload.get("dependencies") or [])
        if isinstance(d, str) and d.strip()
    ]
    requirements_target = application_root / "requirements.txt"
    requirements_target.write_text(
        "\n".join(dependencies) + ("\n" if dependencies else ""),
        encoding="utf-8",
    )
    written_files.append(str(requirements_target.relative_to(final_tmp_dir.parent)))

    # .env.example
    env_vars = payload.get("environment_variables") or []
    if isinstance(env_vars, list) and env_vars:
        env_lines: list[str] = []
        for entry in env_vars:
            if not isinstance(entry, dict):
                continue
            name = str(entry.get("name") or "").strip()
            purpose = str(entry.get("purpose") or "").strip()
            required = bool(entry.get("required"))
            if not name:
                continue
            marker = "REQUIRED" if required else "optional"
            if purpose:
                env_lines.append(f"# {marker}: {purpose}")
            env_lines.append(f"{name}=")
            env_lines.append("")
        env_target = application_root / ".env.example"
        env_target.write_text("\n".join(env_lines).rstrip() + "\n", encoding="utf-8")
        written_files.append(str(env_target.relative_to(final_tmp_dir.parent)))

    return {
        "modality": "application",
        "application_name": app_name,
        "files_written": written_files,
        "bucket_count": sum(
            1 for v in directory_layout.values() if isinstance(v, list)
        ),
    }


# ---------------------------------------------------------------------------
# REQ-trace continuity
# ---------------------------------------------------------------------------


def _scan_req_mentions(root: Path) -> set[str]:
    """Return every REQ-NNN id mentioned in any text file under `root`."""
    found: set[str] = set()
    if not root.exists():
        return found
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        try:
            text = path.read_text(encoding="utf-8")
        except (OSError, UnicodeDecodeError):
            continue
        for match in _REQ_RE.findall(text):
            found.add(match)
    return found


def _check_req_trace_continuity(
    expected_req_ids: set[str],
    fragments: list[dict[str, Any]],
    final_tmp_dir: Path,
    allow_req_drop: tuple[str, ...],
) -> dict[str, Any]:
    """Return a diff payload + raise on forbidden drops."""
    req_in_fragments: set[str] = set()
    for fragment in fragments:
        for req in fragment.get("req_mentions") or []:
            if req in expected_req_ids:
                req_in_fragments.add(req)

    req_in_final = _scan_req_mentions(final_tmp_dir) & expected_req_ids

    synthesis_drops = sorted(req_in_fragments - req_in_final)
    allowed_drops = sorted(set(allow_req_drop) & set(synthesis_drops))
    forbidden_drops = sorted(set(synthesis_drops) - set(allow_req_drop))

    if forbidden_drops:
        sample_locations: dict[str, list[str]] = {}
        for fragment in fragments:
            for req in fragment.get("req_mentions") or []:
                if req in forbidden_drops:
                    locations = sample_locations.setdefault(req, [])
                    if len(locations) < 3:
                        locations.append(_fragment_token(fragment))
        details = "\n  - ".join(
            f"{req} (was in {sample_locations.get(req, [])})"
            for req in forbidden_drops
        )
        raise ValueError(
            "final_synthesis: synthesis dropped REQ-NNN annotations that were "
            "present in work fragments. Re-run the synthesizer with these REQs "
            "preserved, or pass `--allow-req-drop "
            f"{','.join(forbidden_drops)}` to acknowledge the drop:\n  - "
            + details
        )

    return {
        "req_in_fragments": sorted(req_in_fragments),
        "req_in_final": sorted(req_in_final),
        "synthesis_drops": synthesis_drops,
        "allowed_drops": allowed_drops,
    }


# ---------------------------------------------------------------------------
# Postflight entry point
# ---------------------------------------------------------------------------


def run_final_synthesize_finalize(
    artifacts_root: Path,
    workspace_root: Path,
    decision_log_version: int | None = None,
    *,
    allow_req_drop: tuple[str, ...] = (),
    force: bool = False,
) -> dict[str, Any]:
    """Validate subagent returns, materialize executions/v{N}/final/<bucket>/.

    `allow_req_drop` is the operator-supplied list of REQ-NNN ids that may
    be missing from the assembled tree. Drops not on this list cause a
    `ValueError` and `final/` is left untouched.
    """
    paths = build_paths(artifacts_root)
    ensure_layout(paths)

    if not paths.final_synthesis_work_plan_path.exists():
        raise FileNotFoundError(
            f"work plan missing: {paths.final_synthesis_work_plan_path}. "
            "Run `meta-compiler final-synthesize-start` first."
        )
    plan_payload = load_yaml(paths.final_synthesis_work_plan_path) or {}
    work_plan = plan_payload.get("final_synthesis_work_plan")
    if not isinstance(work_plan, dict):
        raise ValueError(
            f"work plan {paths.final_synthesis_work_plan_path} missing root key "
            "'final_synthesis_work_plan'"
        )

    version = int(work_plan.get("decision_log_version") or 0)
    if decision_log_version is not None and decision_log_version != version:
        raise ValueError(
            f"--decision-log-version={decision_log_version} doesn't match "
            f"work plan version v{version}; re-run final-synthesize-start."
        )
    if version <= 0:
        raise ValueError("work plan missing decision_log_version")

    project_type = str(work_plan.get("project_type") or "")
    project_slug = str(work_plan.get("project_slug") or "synthesis")
    work_dir = Path(str(work_plan.get("work_dir") or paths.executions_dir / f"v{version}" / "work"))
    final_dir = paths.final_dir_for(version)
    modality_keys = list(work_plan.get("modality_keys") or [])

    # Load all subagent returns
    returns: dict[str, dict[str, Any]] = {}
    missing: list[str] = []
    for modality in modality_keys:
        payload = _load_subagent_return(paths, modality)
        if payload is None:
            missing.append(modality)
            continue
        returns[modality] = payload

    if missing:
        raise FileNotFoundError(
            "subagent returns missing for modalities: "
            f"{missing}. Invoke .github/prompts/final-synthesis.prompt.md "
            f"to fan out and persist returns under "
            f"{paths.final_synthesis_subagent_returns_dir}."
        )

    # Validate every return; aggregate failures.
    issues_per_modality = _validate_returns(work_plan, returns)
    aggregated: list[str] = []
    for modality, issues in issues_per_modality.items():
        for issue in issues:
            aggregated.append(f"[{modality}] {issue}")
    if aggregated:
        raise ValueError(
            "final-synthesis returns failed validation:\n  - "
            + "\n  - ".join(aggregated)
        )

    # Mtime guard: refuse to overwrite a final/ tree edited after the last
    # report, unless --force.
    report_path = paths.final_synthesis_report_path(version)
    if final_dir.exists() and report_path.exists() and not force:
        report_mtime = report_path.stat().st_mtime
        edited: list[str] = []
        for path in final_dir.rglob("*"):
            if not path.is_file():
                continue
            try:
                if path.stat().st_mtime > report_mtime + 1.0:
                    edited.append(str(path.relative_to(final_dir)))
            except OSError:
                continue
        if edited:
            raise RuntimeError(
                "final/ contains files edited after the last "
                "final_synthesis_report.yaml. Re-running synthesis would "
                "overwrite them. Re-run with --force to override.\n"
                f"  - {edited[:5]}"
                + (f" ... (+{len(edited) - 5} more)" if len(edited) > 5 else "")
            )

    # Apply each modality into a tmp dir; only swap once everything succeeds.
    final_tmp_dir = final_dir.with_suffix(".tmp")
    _ensure_clean_tmp(final_tmp_dir)

    apply_results: dict[str, Any] = {}
    try:
        for modality in modality_keys:
            payload = returns[modality]
            work_plan_slice = (work_plan.get("modalities") or {}).get(modality, {})
            if modality == "library":
                apply_results[modality] = _apply_library_proposal(
                    payload, work_plan_slice, work_dir, final_tmp_dir
                )
            elif modality == "document":
                apply_results[modality] = _apply_document_proposal(
                    payload, work_plan_slice, work_dir, final_tmp_dir, project_slug
                )
            elif modality == "application":
                apply_results[modality] = _apply_application_proposal(
                    payload, work_plan_slice, work_dir, final_tmp_dir
                )
            else:
                raise ValueError(f"unknown modality {modality!r} in work plan")

        # REQ-trace continuity
        all_fragments: list[dict[str, Any]] = []
        for modality in modality_keys:
            slice_fragments = (work_plan.get("modalities") or {}).get(modality, {}).get(
                "fragments"
            ) or []
            all_fragments.extend(slice_fragments)
        # Dedupe by token
        seen: set[str] = set()
        deduped_fragments: list[dict[str, Any]] = []
        for fragment in all_fragments:
            token = _fragment_token(fragment)
            if token in seen:
                continue
            seen.add(token)
            deduped_fragments.append(fragment)

        expected_req_ids = set(work_plan.get("expected_req_ids") or [])
        req_trace_diff = _check_req_trace_continuity(
            expected_req_ids,
            deduped_fragments,
            final_tmp_dir,
            tuple(allow_req_drop),
        )
    except Exception:
        # Discard partial assembly; final/ remains untouched.
        if final_tmp_dir.exists():
            shutil.rmtree(final_tmp_dir, ignore_errors=True)
        raise

    # All modalities applied + REQ check passed — swap into place.
    _swap_tmp_into_place(final_tmp_dir, final_dir)

    # Write the report
    writes: list[str] = []
    for modality_result in apply_results.values():
        writes.extend(modality_result.get("files_written") or [])

    report_payload = {
        "final_synthesis_report": {
            "generated_at": iso_now(),
            "decision_log_version": version,
            "project_type": project_type,
            "project_slug": project_slug,
            "modality_keys": modality_keys,
            "modalities": apply_results,
            "writes": writes,
            "writes_count": len(writes),
            "req_trace_diff": req_trace_diff,
            "allowed_req_drops": [
                {"req_id": req, "reason": "operator override"}
                for req in req_trace_diff.get("allowed_drops") or []
            ],
            "work_plan_path": str(paths.final_synthesis_work_plan_path),
            "final_dir": str(final_dir),
        }
    }
    dump_yaml(report_path, report_payload)

    _set_manifest_stage(paths, "4-synthesized")

    return {
        "status": "synthesized",
        "decision_log_version": version,
        "project_type": project_type,
        "modality_keys": modality_keys,
        "modalities_applied": list(apply_results.keys()),
        "writes_count": len(writes),
        "req_trace_diff": req_trace_diff,
        "report_path": str(report_path),
        "final_dir": str(final_dir),
    }


__all__ = [
    "run_final_synthesize_start",
    "run_final_synthesize_finalize",
]
